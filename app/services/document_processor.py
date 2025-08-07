import io
import re
from typing import List, Dict, Any
from urllib.parse import urlparse
import httpx
from loguru import logger

# Document processing libraries
import PyPDF2
from docx import Document
import pdfplumber
from bs4 import BeautifulSoup

from app.core.config import settings
from app.models.document import DocumentChunk
from app.services.cache_service import cache_service

class DocumentProcessor:
    """Enhanced document processor optimized for insurance policy accuracy"""
    
    def __init__(self):
        # Optimized chunking for better context preservation
        self.chunk_size = 1000       # Larger chunks for more context
        self.chunk_overlap = 200     # More overlap for continuity
        self.max_file_size = settings.max_file_size_bytes
        
        # Enhanced chunking parameters for insurance documents
        self.sentence_endings = r'[.!?]\s+'
        self.paragraph_separators = r'\n\s*\n'
        
        # Insurance-specific patterns for better chunking
        self.policy_section_patterns = [
            r'(?i)(section|clause|article|part|chapter)\s+\d+',
            r'(?i)(coverage|benefit|exclusion|condition|definition)',
            r'(?i)(waiting period|grace period|premium|deductible)',
            r'(?i)(sum insured|policy term|renewal)'
        ]
        
    async def download_document(self, url: str) -> bytes:
        """Download document from URL with enhanced error handling"""
        logger.info(f"Downloading document from: {url}")
        
        # Check cache first
        cache_key = f"download:{url}"
        cached_content = await cache_service.get(cache_key)
        if cached_content:
            logger.info(f"Retrieved document from cache: {url}")
            return cached_content
        
        async with httpx.AsyncClient(timeout=settings.REQUEST_TIMEOUT) as client:
            response = await client.get(url)
            response.raise_for_status()
            
            # Check file size
            content_length = int(response.headers.get('content-length', 0))
            if content_length > self.max_file_size:
                raise ValueError(f"File size {content_length} exceeds maximum allowed size")
            
            content = response.content
            
            # Cache the downloaded content for longer (2 hours for accuracy)
            await cache_service.set(cache_key, content, ttl=7200)
            logger.info(f"Cached downloaded document: {url}")
            
            return content
    
    def extract_text_from_pdf(self, content: bytes) -> str:
        """Enhanced PDF text extraction with better formatting"""
        text = ""
        
        try:
            # Use pdfplumber for better text extraction
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        # Enhanced text cleaning for insurance documents
                        page_text = self._clean_insurance_text(page_text)
                        text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                        
        except Exception as e:
            logger.warning(f"pdfplumber failed, trying PyPDF2: {e}")
            
            # Fallback to PyPDF2
            try:
                pdf_file = io.BytesIO(content)
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    if page_text:
                        page_text = self._clean_insurance_text(page_text)
                        text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                        
            except Exception as e2:
                logger.error(f"PDF extraction failed: {e2}")
                raise
        
        return text.strip()
    
    def _clean_insurance_text(self, text: str) -> str:
        """Enhanced text cleaning for insurance policy documents"""
        # Fix common PDF extraction issues
        text = re.sub(r'\s+', ' ', text)  # Multiple spaces to single
        
        # Fix broken words across lines
        text = re.sub(r'(\w)-\s*\n\s*(\w)', r'\1\2', text)
        
        # Preserve important formatting for policy documents
        text = re.sub(r'(\d+)\s*\.\s*', r'\1. ', text)  # Fix numbered lists
        text = re.sub(r'([a-z])([A-Z])', r'\1. \2', text)  # Add periods between sentences
        
        # Preserve policy-specific formatting
        text = re.sub(r'(?i)(section|clause|article)\s*(\d+)', r'\1 \2:', text)
        text = re.sub(r'(?i)(waiting period|grace period):\s*', r'\1: ', text)
        
        # Clean up excessive newlines but preserve paragraph structure
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r'\n+', ' ', text)  # Convert to single space for better chunking
        
        return text.strip()
    
    def extract_text_from_docx(self, content: bytes) -> str:
        """Enhanced DOCX text extraction"""
        doc = Document(io.BytesIO(content))
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                # Preserve paragraph structure
                para_text = paragraph.text.strip()
                # Check if it looks like a heading
                if len(para_text) < 100 and (para_text.isupper() or para_text.istitle()):
                    text_parts.append(f"\n{para_text}\n")
                else:
                    text_parts.append(para_text)
        
        # Extract from tables with better formatting
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    table_text.append(' | '.join(row_text))
            
            if table_text:
                text_parts.append('\n--- TABLE ---\n' + '\n'.join(table_text) + '\n--- END TABLE ---\n')
        
        return '\n'.join(text_parts)
    
    def extract_text_from_html(self, content: bytes) -> str:
        """Enhanced HTML/Email extraction"""
        soup = BeautifulSoup(content, 'html.parser')
        
        # Remove unwanted elements
        for element in soup(["script", "style", "nav", "header", "footer"]):
            element.decompose()
        
        # Get text with better formatting preservation
        text = soup.get_text(separator='\n', strip=True)
        
        # Clean up while preserving structure
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r'[ \t]+', ' ', text)  # Clean horizontal whitespace
        
        return text
    
    def enhanced_chunk_text(self, text: str, metadata: Dict[str, Any] = None) -> List[DocumentChunk]:
        """Enhanced chunking strategy for insurance policy documents"""
        if metadata is None:
            metadata = {}
        
        chunks = []
        
        # First, identify policy sections for better chunking
        sections = self._identify_policy_sections(text)
        
        if sections:
            # Process each section separately
            for section_name, section_text in sections.items():
                section_chunks = self._chunk_section(section_text, metadata, section_name)
                chunks.extend(section_chunks)
        else:
            # Fallback to paragraph-based chunking
            chunks = self._chunk_by_paragraphs(text, metadata)
        
        logger.info(f"Enhanced chunking created {len(chunks)} chunks")
        return chunks
    
    def _identify_policy_sections(self, text: str) -> Dict[str, str]:
        """Identify different sections in insurance policy"""
        sections = {}
        
        # Common insurance policy section patterns
        section_patterns = [
            (r'(?i)(definitions?|meaning of terms)', 'definitions'),
            (r'(?i)(coverage|benefits? covered)', 'coverage'),
            (r'(?i)(exclusions?|what is not covered)', 'exclusions'),
            (r'(?i)(waiting periods?|grace periods?)', 'waiting_periods'),
            (r'(?i)(claims? procedure|how to claim)', 'claims'),
            (r'(?i)(renewal|continuation)', 'renewal'),
            (r'(?i)(premium|payment)', 'premium'),
            (r'(?i)(conditions|terms)', 'conditions')
        ]
        
        # Try to split text by sections
        current_section = 'general'
        current_text = []
        
        lines = text.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # Check if line matches any section pattern
            section_found = False
            for pattern, section_key in section_patterns:
                if re.search(pattern, line):
                    # Save previous section
                    if current_text:
                        sections[current_section] = '\n'.join(current_text)
                    
                    # Start new section
                    current_section = section_key
                    current_text = [line]
                    section_found = True
                    break
            
            if not section_found:
                current_text.append(line)
        
        # Add final section
        if current_text:
            sections[current_section] = '\n'.join(current_text)
        
        # Only return sections if we found meaningful splits
        if len(sections) > 1:
            return sections
        else:
            return {}
    
    def _chunk_section(self, text: str, metadata: Dict[str, Any], section_name: str) -> List[DocumentChunk]:
        """Chunk a specific policy section"""
        chunks = []
        
        # Split by paragraphs first
        paragraphs = re.split(r'\n\s*\n', text)
        
        current_chunk = ""
        current_word_count = 0
        
        for paragraph in paragraphs:
            paragraph = paragraph.strip()
            if not paragraph:
                continue
            
            paragraph_words = len(paragraph.split())
            
            # If adding this paragraph exceeds chunk size, save current chunk
            if current_word_count + paragraph_words > self.chunk_size and current_chunk:
                chunks.append(self._create_enhanced_chunk(
                    current_chunk, metadata, len(chunks), section_name
                ))
                
                # Start new chunk with overlap
                overlap_text = self._get_overlap_text(current_chunk)
                current_chunk = overlap_text + paragraph
                current_word_count = len(current_chunk.split())
            else:
                current_chunk += ("\n\n" if current_chunk else "") + paragraph
                current_word_count += paragraph_words
        
        # Add final chunk
        if current_chunk:
            chunks.append(self._create_enhanced_chunk(
                current_chunk, metadata, len(chunks), section_name
            ))
        
        return chunks
    
    def _chunk_by_paragraphs(self, text: str, metadata: Dict[str, Any]) -> List[DocumentChunk]:
        """Fallback chunking by paragraphs"""
        chunks = []
        paragraphs = re.split(r'\n\s*\n', text)
        
        current_chunk = ""
        current_word_count = 0
        
        for paragraph in paragraphs:
            paragraph = paragraph.strip()
            if not paragraph:
                continue
            
            paragraph_words = len(paragraph.split())
            
            if current_word_count + paragraph_words > self.chunk_size and current_chunk:
                chunks.append(self._create_enhanced_chunk(
                    current_chunk, metadata, len(chunks)
                ))
                
                # Start new chunk with overlap
                overlap_text = self._get_overlap_text(current_chunk)
                current_chunk = overlap_text + paragraph
                current_word_count = len(current_chunk.split())
            else:
                current_chunk += ("\n\n" if current_chunk else "") + paragraph
                current_word_count += paragraph_words
        
        if current_chunk:
            chunks.append(self._create_enhanced_chunk(
                current_chunk, metadata, len(chunks)
            ))
        
        return chunks
    
    def _get_overlap_text(self, text: str) -> str:
        """Get overlap text for chunk continuity"""
        words = text.split()
        if len(words) <= self.chunk_overlap:
            return text + " "
        
        # Take last chunk_overlap words
        overlap_words = words[-self.chunk_overlap:]
        return " ".join(overlap_words) + " "
    
    def _create_enhanced_chunk(self, content: str, metadata: Dict[str, Any], index: int, section: str = None) -> DocumentChunk:
        """Create an enhanced document chunk with better metadata"""
        enhanced_metadata = {
            **metadata,
            'chunk_index': index,
            'word_count': len(content.split()),
            'char_count': len(content),
            'section': section or 'general',
            'contains_numbers': bool(re.search(r'\d+', content)),
            'contains_percentages': bool(re.search(r'\d+%', content)),
            'contains_dates': bool(re.search(r'\d{1,2}[/-]\d{1,2}[/-]\d{2,4}', content)),
            'policy_terms': self._extract_policy_terms(content)
        }
        
        return DocumentChunk(
            content=content.strip(),
            metadata=enhanced_metadata
        )
    
    def _extract_policy_terms(self, content: str) -> List[str]:
        """Extract important policy terms from content"""
        terms = []
        
        # Important insurance terms to identify
        policy_terms = [
            'premium', 'deductible', 'coverage', 'benefit', 'exclusion',
            'waiting period', 'grace period', 'sum insured', 'co-payment',
            'renewal', 'claim', 'policy term', 'maturity', 'rider'
        ]
        
        content_lower = content.lower()
        for term in policy_terms:
            if term in content_lower:
                terms.append(term)
        
        return terms
    
    async def process_document(self, url: str) -> List[DocumentChunk]:
        """Enhanced document processing for maximum accuracy"""
        logger.info(f"Processing document with enhanced accuracy: {url}")
        
        # Check if chunks are already cached
        cached_chunks = await cache_service.get_document_chunks(url)
        if cached_chunks:
            logger.info(f"Retrieved {len(cached_chunks)} cached chunks for document: {url}")
            return cached_chunks
        
        # Download document
        content = await self.download_document(url)
        
        # Determine file type from URL
        parsed_url = urlparse(url)
        filename = parsed_url.path.split('/')[-1].lower()
        
        # Cache key for extracted text
        content_hash = str(hash(content))
        text_cache_key = f"extracted_text:{filename}:{content_hash}"
        
        # Check if text is already extracted and cached
        cached_text_data = await cache_service.get(text_cache_key)
        if cached_text_data:
            text = cached_text_data.get('text', '')
            doc_type = cached_text_data.get('doc_type', 'unknown')
            logger.info(f"Retrieved cached extracted text for: {filename}")
        else:
            # Extract text based on file type with enhanced methods
            if filename.endswith('.pdf'):
                text = self.extract_text_from_pdf(content)
                doc_type = 'pdf'
            elif filename.endswith('.docx'):
                text = self.extract_text_from_docx(content)
                doc_type = 'docx'
            elif filename.endswith('.html') or filename.endswith('.eml'):
                text = self.extract_text_from_html(content)
                doc_type = 'email'
            else:
                # Try to decode as plain text
                text = content.decode('utf-8', errors='ignore')
                doc_type = 'text'
            
            # Cache extracted text for longer (4 hours for accuracy)
            text_data = {'text': text, 'doc_type': doc_type}
            await cache_service.set(text_cache_key, text_data, ttl=14400)
            logger.info(f"Cached extracted text for: {filename}")
        
        if not text or len(text.strip()) < 50:
            raise ValueError("Document appears to be empty or too short")
        
        # Create enhanced metadata
        metadata = {
            'source_url': url,
            'document_type': doc_type,
            'filename': filename,
            'total_chars': len(text),
            'total_words': len(text.split()),
            'is_insurance_policy': self._detect_insurance_policy(text),
            'language': 'english'  # Could be enhanced with language detection
        }
        
        # Use enhanced chunking
        chunks = self.enhanced_chunk_text(text, metadata)
        
        # Cache the chunks for longer (4 hours)
        await cache_service.cache_document_chunks(url, chunks, ttl=14400)
        
        logger.info(f"Created and cached {len(chunks)} enhanced chunks from document")
        return chunks
    
    def _detect_insurance_policy(self, text: str) -> bool:
        """Detect if document is an insurance policy"""
        insurance_keywords = [
            'policy', 'premium', 'coverage', 'benefit', 'exclusion',
            'insured', 'insurer', 'claim', 'deductible', 'sum insured'
        ]
        
        text_lower = text.lower()
        keyword_count = sum(1 for keyword in insurance_keywords if keyword in text_lower)
        
        return keyword_count >= 5  # If 5+ insurance keywords found

# Singleton instance
document_processor = DocumentProcessor()
